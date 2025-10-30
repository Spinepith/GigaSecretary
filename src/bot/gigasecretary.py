from langchain_gigachat import GigaChat
from langchain.prompts.chat import ChatPromptTemplate
from langchain_core.runnables import RunnableSequence
from langchain.schema import HumanMessage, AIMessage
from langchain_core.chat_history import InMemoryChatMessageHistory

import io
import docx

from ..db import db
from ..bot import utils
from ..config import *


llm = GigaChat(credentials=GIGACHAT_TOKEN, verify_ssl_certs=False)
user_memory = {}

main_prompt = '''
    Ты — AI-секретарь (Умный секретарь) в системе электронного документооборота (РСЭД). Твоя задача — профессионально консультировать по вопросам документооборота в компании.

    Тебе доступен список шаблонов документов: {templates}
    Тебе доступен список отделов компании: {departments}
    
    АЛГОРИТМ РАБОТЫ:
    1. Проанализируй запрос пользователя и определи, можно ли его интерпретировать в контексте документооборота РСЭД
    2. Если запрос можно связать с документооборотом - дай развернутый ответ в этом контексте
    3. Если запрос НЕЛЬЗЯ связать с документооборотом РСЭД - вежливо откажись отвечать
    
    КРИТЕРИИ ИНТЕРПРЕТАЦИИ:
    - Вопросы о шаблонах документов → можно интерпретировать как запрос о подходящем шаблоне из списка
    - Вопросы о направлениях документов → можно интерпретировать как консультацию по отделам из списка
    - Вопросы об оформлении документов → можно интерпретировать как помощь в документообороте РСЭД
    - Вопросы о процедурах документооборота → прямо касаются рабочих процессов компании
    
    ПРИМЕРЫ КОРРЕКТНОЙ ИНТЕРПРЕТАЦИИ:
    - "Какой шаблон использовать для служебной записки?" → "Для служебной записки в РСЭД рекомендую использовать шаблон СЛ-2024..."
    - "В какой отдел направлять заявку на оборудование?" → "Заявки на оборудование в РСЭД направляются в отдел технического обеспечения..."
    - "Как оформить командировочное удостоверение?" → "Командировочное удостоверение в РСЭД оформляется по форме КУ-стандарт..."
    
    ПРИМЕРЫ ОТКАЗА:
    - "Какая погода завтра?" → нельзя интерпретировать в контексте документооборота
    - "Расскажи анекдот" → нельзя интерпретировать в контексте документооборота
    - "Как приготовить ужин?" → нельзя интерпретировать в контексте документооборота
    
    ПРАВИЛА ОТВЕТА:
    - Отвечай сразу по сути, без вводных фраз
    - Не пиши "Анализ запроса", "Интерпретация" и т.д.
    - Не объясняй ход своих мыслей
    - Только конкретный и информативный ответ на вопрос
    - Подбирай шаблоны только из предоставленного списка
    - Рекомендуй отделы только из предоставленного списка
    - Если не можешь найти подходящий шаблон или отдел - ответь: "В доступных данных нет подходящего варианта для вашего запроса"
    - Всегда уточняй, доступен ли подходящий шаблон/отдел в данный момент
    - При необходимости дальнейшей консультации предложи продолжить обсуждение в этом чате
    - Не предлагай конкретных способов связи кроме продолжения диалога в этом интерфейсе
    
    Всегда старайся найти связь с документооборотом РСЭД, если это возможно. Отказывай только если связь действительно отсутствует.
    
    История диалога: {chat_history}
    Пользователь: {user_input}
    AI-секретарь:
'''


def ask(user_id: int, user_input: str):
    prompt = ChatPromptTemplate.from_template(main_prompt)
    return __get_response(prompt, user_id, user_input)


def analyze_document(user_id: int, user_input: str):
    analyze_prompt = '''
        ДОПОЛНИТЕЛЬНЫЕ ИНСТРУКЦИИ ДЛЯ АНАЛИЗА ДОКУМЕНТА:

        АЛГОРИТМ АНАЛИЗА ДОКУМЕНТА:
        1. Определи тип документа на основе содержания
        2. Составь краткое описание документа (3-4 предложения)
        3. Определи подходящий отдел для направления из доступного списка
        4. Сформулируй ответ в строгом формате
        
        ДОПУСТИМЫЕ ТИПЫ ДОКУМЕНТОВ:
        - Заявление
        - Служебная записка  
        - Приказ
        - Договор
        - Счет
        - Жалоба
        - Запрос
        - Отчет
        - Официальные заявления, служебные записки, запросы
        - Тексты, содержащие реквизиты, официальные обращения
        - Деловые предложения, коммерческие документы
        - Жалобы, претензии, официальные обращения
        - Любые тексты, относящиеся к рабочим процессам компании
        - Другое (если не подходит под категории)
        
        НЕ СЧИТАТЬ ДОКУМЕНТОМ:
        - Приветствия, прощания, личные сообщения
        - Вопросы о погоде, личные просьбы не по теме
        - Спам, реклама, бессвязный текст
        - Пустые сообщения или сообщения без смысловой нагрузки
        - Просьбы с абсурдной просьбой связанной с документами
        
        ФОРМАТ ОТВЕТА ПРИ АНАЛИЗЕ:
        Тип документа: [определенный тип]
        Описание: [краткое описание содержания]
        Рекомендуемый отдел: [конкретный отдел из списка]
        [Отдел: название_отдела]
        
        КРИТИЧЕСКИЕ ПРАВИЛА:
        - Всегда заканчивай ответ конструкцией [Отдел: название_отдела]
        - Если отдел не определен - [Отдел: Неизвестно]
        - Не добавляй текст после указания отдела
        - Используй только отделы из предоставленного списка
        - Описание должно быть объективным на основе текста документа
    '''
    prompt = ChatPromptTemplate.from_template(main_prompt + analyze_prompt)
    return __get_response(prompt, user_id, user_input)


def compare_documents(user_id: int, user_input: str, template: str):
    compare_prompt = """
        ДОПОЛНИТЕЛЬНЫЕ ИНСТРУКЦИИ ДЛЯ СРАВНЕНИЯ ДОКУМЕНТОВ:

        АЛГОРИТМ СРАВНЕНИЯ:
        1. Определи тип документа и отдел для предоставленного документа
        2. Проведи сравнительный анализ с шаблоном по критериям соответствия
        3. Выяви все несоответствия
        4. Сформулируй рекомендации по исправлению
        
        КРИТЕРИИ СРАВНЕНИЯ:
        - Структура документа: наличие обязательных разделов, порядок разделов
        - Содержание: ключевые реквизиты, полнота информации, соответствие цели  
        - Оформление: формальность тона, стандартные формулировки, форматирование
        
        ФОРМАТ ОТВЕТА ПРИ СРАВНЕНИИ:
        Тип документа: [определенный тип]
        Рекомендуемый отдел: [конкретный отдел из списка]
        
        Несоответствия шаблону:
        1. [конкретное несоответствие 1]
        2. [конкретное несоответствие 2]
        ...
        
        Рекомендации:
        1. [рекомендация по исправлению 1]  
        2. [рекомендация по исправлению 2]
        ...
        
        [Отдел: название_отдела]
        
        КРИТИЧЕСКИЕ ПРАВИЛА:
        - Всегда заканчивай ответ конструкцией [Отдел: название_отдела]
        - Если отдел не определен - [Отдел: Неизвестно]
        - Не добавляй текст после указания отдела
        - Несоответствия формулируй конкретно и объективно
        - Рекомендации должны быть практичными и выполнимыми
        {template}
    """

    prompt = ChatPromptTemplate.from_template(main_prompt + compare_prompt)
    return __get_response(prompt, user_id, user_input, template=template)


# ПЛОХО РАБОТАЕТ, НАДО РАБОТАТЬ НАД ЭТИМ [НЕПРАВИЛЬНО ЗАПОЛНЯЕТ СТИЛИ И ПО СМЫСЛУ НЕ ПРАВИЛЬНО ВСТАВЛЯЕТ]
def fill_document(user_id: int, user_input: str, template: str):
    try:
        try:
            doc = docx.Document(template)
        except FileNotFoundError:
            utils.log_file("Ошибка при подходе заполнить файл: ФАЙЛ НЕ НАЙДЕН")
            return

        line_analysis_prompt = """
            ДОПОЛНИТЕЛЬНЫЕ ИНСТРУКЦИИ ДЛЯ АНАЛИЗА СТРОКИ ОФИЦИАЛЬНОГО ДОКУМЕНТА:
            
            АНАЛИЗИРУЙ КАЖДУЮ СТРОКУ ДОКУМЕНТА:
            
            ШАГ АНАЛИЗА:
            1. Определи тип данных, которые должны быть в этой строке (ФИО, дата, сумма и т.д.)
            2. Найди в пользовательском вводе соответствующие данные этого типа
            3. Если нашел - заполни строку, заменив пропуски на найденные данные
            4. Если не нашел - оставь строку без изменений
            
            ТИПЫ ДАННЫХ ДЛЯ ЗАПОЛНЕНИЯ:
            - ФИО: полное имя человека (Иванов Иван Иванович)
            - Дата: конкретная дата в любом формате
            - Адрес: место проживания или нахождения
            - Сумма: числовое значение с валютами или без
            - Название: наименование организации, товара, услуги
            - Номер: цифровые идентификаторы, телефоны
            - Описание: детальное описание чего-либо
            
            ФОРМАТ ОТВЕТА:
            [заполненная_строка]
            
            СТРОГИЕ ПРАВИЛА:
            - Заполняй ТОЛЬКО если точно нашел соответствующие данные в пользовательском вводе
            - Сохраняй ВСЕ форматирование и структуру исходной строки
            - Заменяй ТОЛЬКО пропуски (___, __________), оставляя остальной текст неизменным
            - Если не нашел подходящих данных - верни исходную строку без изменений
            - Не добавляй новые предложения или абзацы
            - Не изменяй текст, который уже есть в строке (кроме пропусков)
        """

        for paragraph in doc.paragraphs:
            original_text = paragraph.text.strip()

            if not original_text:
                continue

            has_placeholders = ('___' in original_text or '__________' in original_text)

            if has_placeholders:
                analysis_prompt = f"""
                    СТРОКА ДОКУМЕНТА: "{original_text}"
                    ДАННЫЕ ПОЛЬЗОВАТЕЛЯ: "{user_input}"
                    
                    Проанализируй, какие данные должны быть на месте пропусков в этой строке.
                    Найди в пользовательском вводе ПОДХОДЯЩИЕ данные для заполнения пропусков.
                    Если нашел - заполни пропуски, сохраняя весь остальной текст без изменений.
                    Если не нашел - оставь строку как есть.
                """
                full_prompt = ChatPromptTemplate.from_template(main_prompt + line_analysis_prompt + analysis_prompt)
                response = __get_response(full_prompt, user_id, user_input)

                if response.startswith('[') and response.endswith(']'):
                    filled_text = response[1:-1].strip()
                else:
                    filled_text = response.strip()

                if (filled_text and
                        filled_text != original_text and
                        not any(nonsense in filled_text for nonsense in ['Хпзрзухпзпз', 'щащузу', 'ащвщ', 'лалада', 'Отрицаю'])):

                    if paragraph.runs:
                        first_run = paragraph.runs[0]
                        original_style = first_run.style
                        paragraph.clear()
                        new_run = paragraph.add_run(filled_text)
                        new_run.style = original_style
                    else:
                        paragraph.text = filled_text

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text.strip()

                        if not original_text:
                            continue

                        has_placeholders = ('___' in original_text or '__________' in original_text)

                        if has_placeholders:
                            analysis_prompt = f"""
                                СТРОКА В ТАБЛИЦЕ: "{original_text}"
                                ДАННЫЕ ПОЛЬЗОВАТЕЛЯ: "{user_input}"
                                
                                Проанализируй, какие данные должны быть на месте пропусков в этой строке таблицы.
                                Найди в пользовательском вводе ПОДХОДЯЩИЕ данные для заполнения пропусков.
                                Если нашел - заполни пропуски, сохраняя весь остальной текст без изменений.
                                Если не нашел - оставь строку как есть.
                            """
                            full_prompt = ChatPromptTemplate.from_template(
                                main_prompt + line_analysis_prompt + analysis_prompt
                            )
                            response = __get_response(full_prompt, user_id, user_input)

                            if response.startswith('[') and response.endswith(']'):
                                filled_text = response[1:-1].strip()
                            else:
                                filled_text = response.strip()

                            if (filled_text and filled_text != original_text and
                                    not any(nonsense in filled_text for nonsense in
                                            ['Хпзрзухпзпз', 'щащузу', 'ащвщ', 'лалада', 'Отрицаю'])):

                                if paragraph.runs:
                                    first_run = paragraph.runs[0]
                                    original_style = first_run.style
                                    paragraph.clear()
                                    new_run = paragraph.add_run(filled_text)
                                    new_run.style = original_style
                                else:
                                    paragraph.text = filled_text

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    except Exception as e:
        utils.log_file(f"Ошибка при попытке заполнить документ: {str(e)}")


def __get_response(prompt: ChatPromptTemplate, user_id: int, user_input: str, **kwargs):
    main_chain = RunnableSequence(prompt | llm)

    if user_id not in user_memory:
        user_memory[user_id] = InMemoryChatMessageHistory()

    memory = user_memory[user_id]

    chat_history = ""
    for message in memory.messages:
        if isinstance(message, HumanMessage):
            chat_history += f"Пользователь: {message.content}\n"
        elif isinstance(message, AIMessage):
            chat_history += f"AI-секретарь: {message.content}\n"

    variables = {
        "user_input": user_input,
        "chat_history": chat_history,
        "departments": db.get_departments(),
        "templates": utils.get_templates(),
        **kwargs
    }

    response_obj = main_chain.invoke(variables)
    response_text = response_obj.content

    memory.add_message(HumanMessage(content=user_input))
    memory.add_message(AIMessage(content=response_text))

    return response_text
